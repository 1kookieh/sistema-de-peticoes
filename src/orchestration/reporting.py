"""Relatórios estruturados de conformidade formal."""
from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from jinja2 import Environment, FileSystemLoader, select_autoescape

from src.core.profiles import ValidationProfile, get_profile
from src.core.validation.docx import _cm, validar

TEMPLATES_DIR = Path(__file__).resolve().parents[2] / "templates"


def _round_cm(value) -> float:
    return round(_cm(value), 2)


def extract_docx_structure(path: Path, profile_id: str | None = None) -> dict[str, Any]:
    profile = get_profile(profile_id)
    doc = Document(str(path))
    sec = doc.sections[0]
    paragraphs = [(p, p.text.strip()) for p in doc.paragraphs]
    texts = [text for _, text in paragraphs if text]
    first_index = next((i for i, (_, text) in enumerate(paragraphs) if text), None)

    blank_lines = 0
    if first_index is not None:
        for _, text in paragraphs[first_index + 1 :]:
            if text:
                break
            blank_lines += 1

    font_names = sorted({
        run.font.name
        for p, _ in paragraphs
        for run in p.runs
        if run.font.name
    })
    font_sizes = sorted({
        float(run.font.size.pt)
        for p, _ in paragraphs
        for run in p.runs
        if run.font.size
    })
    upper_text = "\n".join(texts).upper()
    sections_found = [
        section
        for section in profile.required_sections
        if section.upper() in upper_text
    ]

    return {
        "profile_id": profile.id,
        "page": {
            "width_cm": _round_cm(sec.page_width),
            "height_cm": _round_cm(sec.page_height),
        },
        "margins_cm": {
            "top": _round_cm(sec.top_margin),
            "left": _round_cm(sec.left_margin),
            "bottom": _round_cm(sec.bottom_margin),
            "right": _round_cm(sec.right_margin),
        },
        "paragraph_count": len(doc.paragraphs),
        "first_non_empty": texts[0] if texts else "",
        "blank_lines_after_header": blank_lines,
        "font_names": font_names,
        "font_sizes": font_sizes,
        "contains_oab": "OAB" in upper_text,
        "contains_local_data_hint": " DE 20" in upper_text,
        "required_sections_found": sections_found,
    }


def build_docx_report(
    path: Path,
    profile_id: str | None = None,
    *,
    problems: list[str] | None = None,
) -> dict[str, Any]:
    problemas = validar(path, profile_id=profile_id) if problems is None else problems
    return {
        "path": str(path),
        "status": "ok" if not problemas else "invalid_docx",
        "generated_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "problems": problemas,
        "structure": extract_docx_structure(path, profile_id),
    }


def build_run_report(
    *,
    profile: ValidationProfile,
    strict: bool,
    no_outbox: bool,
    summary: dict[str, int],
    items: list[dict[str, Any]],
) -> dict[str, Any]:
    return {
        "generated_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "profile": {
            "id": profile.id,
            "descricao": profile.descricao,
        },
        "strict": strict,
        "no_outbox": no_outbox,
        "summary": summary,
        "items": items,
    }


def write_json_report(path: Path, report: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")


def render_report_html(report: dict[str, Any]) -> str:
    """Converte um relatório JSON em HTML local para revisão humana."""
    env = Environment(
        loader=FileSystemLoader(TEMPLATES_DIR),
        autoescape=select_autoescape(("html", "xml")),
    )
    template = env.get_template("report.html")
    return template.render(
        generated_at=report.get("generated_at", ""),
        profile=report.get("profile", {}),
        summary=report.get("summary", {}),
        items=report.get("items", []),
    )


def write_html_report(path: Path, report: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(render_report_html(report), encoding="utf-8")


