"""
Converte um texto de petição em .docx seguindo a skill WORD/DOCX JUDICIÁRIO
(padrão forense + base ABNT).

Regras aplicadas:
- A4, margens 3/3/2/2 cm (sup/esq/inf/dir)
- Times New Roman 12, preto
- Justificado, 1,5 entre linhas
- Espaçamento antes/depois: 0 pt
- Recuo de primeira linha: 2,5 cm no corpo do texto
- 7 linhas em branco após o endereçamento da vara
- Negrito apenas em: endereçamento, nome da peça, títulos (DOS FATOS,
  DO DIREITO, etc.), marcadores de alíneas "a)" "b)" "c)", nome do advogado, OAB
- Nome do advogado e OAB centralizados em linhas separadas, sem linha
  para assinatura
- Fechamento "Termos em que, pede deferimento." em parágrafo justificado normal
- Local e data com alinhamento à direita

Uso:
    python -m src.formatar_docx <entrada.txt> <saida.docx>
    python -m src.formatar_docx -           # lê da stdin
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


HEADER_RE = re.compile(
    r"^\s*(EXCELENT[ÍI]SSIMO|AO\s+TABELIONATO|AO\s+INSTITUTO)", re.IGNORECASE
)

TITLE_RE = re.compile(
    r"^\s*(A[ÇC][ÃA]O\s+DE\s+|RECURSO\s+|PETI[ÇC][ÃA]O\s+DE\s+|"
    r"MANIFESTA[ÇC][ÃA]O\s+|QUESITOS\s+PERICIAIS|MANDADO\s+DE\s+SEGURAN[ÇC]A)",
    re.IGNORECASE,
)

SECTION_NAMES = [
    "DOS FATOS", "DO DIREITO", "DA TUTELA DE URGÊNCIA", "DA TUTELA DE URGENCIA",
    "DOS PEDIDOS", "DO VALOR DA CAUSA", "DAS PROVAS",
    "DOS QUESITOS PERICIAIS", "DA COMPLEMENTAÇÃO DOS QUESITOS",
    "DA QUALIDADE DE SEGURADO", "DA INCAPACIDADE LABORAL",
]

SECTION_ROMAN_RE = re.compile(
    r"^\s*([IVX]+)\s*[-–—]\s*(.+)$", re.IGNORECASE
)

ALINEA_RE = re.compile(r"^\s*([a-z](?:\.\d+)?\))\s+(.+)$", re.IGNORECASE)

CLOSING_RE = re.compile(r"^(Termos em que|Nestes termos).*pede deferimento", re.IGNORECASE)
LOCAL_DATA_RE = re.compile(
    r"^[A-ZÁÉÍÓÚÂÊÔÃÕÇ][A-Za-zÁÉÍÓÚÂÊÔÃÕÇáéíóúâêôãõç]+(?:/[A-Z]{2})?,\s*\d{1,2}\s+de\s+\w+\s+de\s+\d{4}"
)
OAB_RE = re.compile(r"^\s*OAB\s*/?\s*[A-Z]{2}\s*[\d.]+\s*$", re.IGNORECASE)


def _setup_page(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(3)
        section.left_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.right_margin = Cm(2)


def _style_run(run) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)


def _new_paragraph(doc: Document, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                   indent: bool = True) -> "object":
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if indent:
        pf.first_line_indent = Cm(2.5)
    p.alignment = align
    return p


def _add_text(p, texto: str, *, bold: bool = False) -> None:
    run = p.add_run(texto)
    run.bold = bold
    _style_run(run)


def _add_simple(doc: Document, texto: str, *, bold: bool = False,
                align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent: bool = False) -> None:
    p = _new_paragraph(doc, align=align, indent=indent)
    _add_text(p, texto, bold=bold)


def _add_blank_lines(doc: Document, n: int) -> None:
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)


def _is_section_title(linha: str) -> bool:
    s = linha.strip().rstrip(".:")
    if not s:
        return False
    if s.upper() in SECTION_NAMES:
        return True
    if SECTION_ROMAN_RE.match(s):
        return True
    if s.isupper() and 3 <= len(s) <= 80 and not HEADER_RE.match(s) and not TITLE_RE.match(s):
        keywords = ("FATO", "DIREITO", "PEDIDO", "PROVA", "CAUSA", "TUTELA",
                    "QUESITO", "QUALIDADE", "INCAPACIDADE", "MÉRITO", "MERITO")
        return any(k in s for k in keywords)
    return False


def _is_title(linha: str) -> bool:
    s = linha.strip()
    if not s.isupper() or len(s) < 10:
        return False
    return bool(TITLE_RE.match(s))


def _is_header(linha: str) -> bool:
    return bool(HEADER_RE.match(linha.strip()))


def _is_closing(linha: str) -> bool:
    return bool(CLOSING_RE.match(linha.strip()))


def _is_local_data(linha: str) -> bool:
    return bool(LOCAL_DATA_RE.match(linha.strip()))


def _is_oab(linha: str) -> bool:
    return bool(OAB_RE.match(linha.strip()))


def _render_alinea(doc: Document, marcador: str, texto: str) -> None:
    p = _new_paragraph(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True)
    _add_text(p, marcador + " ", bold=True)
    _add_text(p, texto, bold=False)


def _render_bloco_justificado(doc: Document, texto: str) -> None:
    p = _new_paragraph(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True)
    _add_text(p, texto, bold=False)


def _tem_alinea(linhas: list[str]) -> bool:
    return any(ALINEA_RE.match(l.strip()) for l in linhas)


def _render_sequencia_alineas(doc: Document, linhas: list[str]) -> None:
    """Cada linha `a) ...`, `b) ...` vira um parágrafo próprio.
    Linhas não-alínea são agregadas à alínea anterior (continuação) ou viram
    parágrafo justificado solto."""
    atual_marcador: str | None = None
    atual_texto: list[str] = []

    def flush():
        nonlocal atual_marcador, atual_texto
        if atual_marcador is not None:
            _render_alinea(doc, atual_marcador, " ".join(atual_texto).strip())
        elif atual_texto:
            _render_bloco_justificado(doc, " ".join(atual_texto).strip())
        atual_marcador = None
        atual_texto = []

    for linha in linhas:
        s = linha.strip()
        m = ALINEA_RE.match(s)
        if m:
            flush()
            atual_marcador = m.group(1)
            atual_texto = [m.group(2).strip()]
        else:
            atual_texto.append(s)
    flush()


def renderizar(texto: str, destino: Path) -> Path:
    doc = Document()
    _setup_page(doc)

    linhas_brutas = texto.splitlines()
    blocos: list[list[str]] = []
    atual: list[str] = []
    for linha in linhas_brutas:
        if linha.strip():
            atual.append(linha.rstrip())
        else:
            if atual:
                blocos.append(atual)
                atual = []
    if atual:
        blocos.append(atual)

    endereco_feito = False
    for bloco in blocos:
        primeira = bloco[0].strip()
        texto_bloco = " ".join(l.strip() for l in bloco)

        if _is_header(primeira) and not endereco_feito:
            for linha in bloco:
                _add_simple(doc, linha.strip(), bold=True,
                            align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
            _add_blank_lines(doc, 7)
            endereco_feito = True
            continue

        if _is_title(primeira) and len(bloco) == 1:
            _add_simple(doc, primeira, bold=True,
                        align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
            continue

        if _is_section_title(primeira):
            _add_simple(doc, primeira, bold=True,
                        align=WD_ALIGN_PARAGRAPH.LEFT, indent=False)
            resto_linhas = bloco[1:]
            if resto_linhas:
                if _tem_alinea(resto_linhas):
                    _render_sequencia_alineas(doc, resto_linhas)
                else:
                    _render_bloco_justificado(
                        doc, " ".join(l.strip() for l in resto_linhas)
                    )
            continue

        if _tem_alinea(bloco):
            _render_sequencia_alineas(doc, bloco)
            continue

        if _is_closing(primeira):
            _render_bloco_justificado(doc, texto_bloco)
            continue

        if _is_local_data(primeira) and len(bloco) == 1:
            _add_simple(doc, primeira, bold=False,
                        align=WD_ALIGN_PARAGRAPH.RIGHT, indent=False)
            continue

        if len(bloco) == 2 and _is_oab(bloco[1].strip()):
            _add_simple(doc, bloco[0].strip(), bold=True,
                        align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
            _add_simple(doc, bloco[1].strip(), bold=True,
                        align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
            continue

        if _is_oab(primeira) and len(bloco) == 1:
            _add_simple(doc, primeira, bold=True,
                        align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
            continue

        _render_bloco_justificado(doc, texto_bloco)

    destino.parent.mkdir(parents=True, exist_ok=True)
    doc.save(destino)
    return destino


def main(argv: list[str]) -> int:
    if len(argv) != 3:
        print(__doc__)
        return 2
    entrada, saida = argv[1], argv[2]
    texto = sys.stdin.read() if entrada == "-" else Path(entrada).read_text(encoding="utf-8")
    destino = renderizar(texto, Path(saida))
    print(f"OK: {destino}")
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv))
