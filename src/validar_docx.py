"""Valida se um .docx atende às regras estruturais do prompt de formatação.

Retorna lista de violações (strings). Uma lista vazia significa que o documento
passou em todas as verificações determinísticas (margens, fonte, alinhamentos,
negrito nos lugares certos, 7 linhas após o endereçamento, OAB no fechamento).

Não valida semântica jurídica — essa responsabilidade fica com o redator
(humano ou integrador externo) ao comparar o texto final contra o
`prompt_peticao.md`.
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

TOP_MARGIN_CM = 3.0
LEFT_MARGIN_CM = 3.0
BOTTOM_MARGIN_CM = 2.0
RIGHT_MARGIN_CM = 2.0
FONTES_ACEITAS = {"Times New Roman", "Arial"}
TAM_FONTE_PT = 12
BLANKS_APOS_ENDERECO = 7

_EMU_POR_CM = 360000


def _cm(emu) -> float:
    return (emu or 0) / _EMU_POR_CM


def validar(path: Path) -> list[str]:
    if not path.exists():
        return [f"arquivo não encontrado: {path}"]

    doc = Document(str(path))
    problemas: list[str] = []

    sec = doc.sections[0]
    for nome, atual, esperado in [
        ("superior", _cm(sec.top_margin), TOP_MARGIN_CM),
        ("esquerda", _cm(sec.left_margin), LEFT_MARGIN_CM),
        ("inferior", _cm(sec.bottom_margin), BOTTOM_MARGIN_CM),
        ("direita", _cm(sec.right_margin), RIGHT_MARGIN_CM),
    ]:
        if abs(atual - esperado) > 0.1:
            problemas.append(
                f"margem {nome}: {atual:.2f} cm (esperado {esperado} cm)"
            )

    fontes_vistas: set[str] = set()
    tamanhos_vistos: set[int] = set()
    for i, p in enumerate(doc.paragraphs):
        for run in p.runs:
            if run.font.name:
                fontes_vistas.add(run.font.name)
            if run.font.size:
                tamanhos_vistos.add(run.font.size.pt)
    fontes_invalidas = fontes_vistas - FONTES_ACEITAS
    if fontes_invalidas:
        problemas.append(
            f"fontes não autorizadas: {sorted(fontes_invalidas)} "
            f"(esperado: {sorted(FONTES_ACEITAS)})"
        )
    tamanhos_invalidos = {t for t in tamanhos_vistos if t != TAM_FONTE_PT}
    if tamanhos_invalidos:
        problemas.append(
            f"tamanhos de fonte não autorizados: {sorted(tamanhos_invalidos)} "
            f"(esperado: {TAM_FONTE_PT}pt)"
        )

    paragrafos = list(doc.paragraphs)
    paragrafos_texto = [(i, p, p.text.strip()) for i, p in enumerate(paragrafos)]
    primeiros_nao_vazios = [(i, p, t) for i, p, t in paragrafos_texto if t]

    if not primeiros_nao_vazios:
        problemas.append("documento vazio")
        return problemas

    idx_p0, p0, t0 = primeiros_nao_vazios[0]
    t0_upper = t0.upper()
    if not t0_upper.startswith(
        ("EXCELENT", "AO TABELIONATO", "AO INSTITUTO", "À ", "AO ")
    ):
        problemas.append(
            f"1º parágrafo não parece endereçamento: {t0[:60]!r}"
        )
    if p0.alignment != WD_ALIGN_PARAGRAPH.CENTER:
        problemas.append("endereçamento não está centralizado")
    if not any(r.bold for r in p0.runs if r.text.strip()):
        problemas.append("endereçamento não está em negrito")

    brancos = 0
    for p in paragrafos[idx_p0 + 1 :]:
        if not p.text.strip():
            brancos += 1
        else:
            break
    if brancos < BLANKS_APOS_ENDERECO:
        problemas.append(
            f"{brancos} linhas em branco após endereçamento "
            f"(esperado {BLANKS_APOS_ENDERECO})"
        )

    ultimos = [t for _, _, t in paragrafos_texto if t][-4:]
    if not any("OAB" in t.upper() for t in ultimos):
        problemas.append("fechamento sem linha de OAB nos últimos parágrafos")

    linha_oab = next(
        (t for _, _, t in reversed(paragrafos_texto) if "OAB" in t.upper()),
        None,
    )
    if linha_oab and len(linha_oab) > 40:
        problemas.append(
            f"linha OAB muito longa (provavelmente nome+OAB juntos): {linha_oab[:60]!r}"
        )

    for _, _, t in paragrafos_texto:
        if "_" * 5 in t or "-" * 10 in t:
            problemas.append(
                "há linha de assinatura (traços/underscores) — proibido"
            )
            break

    return problemas


def main(argv: list[str]) -> int:
    if len(argv) != 2:
        print("uso: python -m src.validar_docx <arquivo.docx>")
        return 2
    problemas = validar(Path(argv[1]))
    if not problemas:
        print("OK")
        return 0
    print(f"VIOLACOES ({len(problemas)}):")
    for v in problemas:
        print(f"  - {v}")
    return 1


if __name__ == "__main__":
    sys.exit(main(sys.argv))
