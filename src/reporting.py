"""Relatórios estruturados de conformidade formal."""
from __future__ import annotations

import json
from html import escape
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document

from src.profiles import ValidationProfile, get_profile
from src.validar_docx import _cm, validar


def _round_cm(value) -> float:
    return round(_cm(value), 2)


def extract_docx_structure(path: Path, profile_id: str | None = None) -> dict[str, Any]:
    profile = get_profile(profile_id)
    doc = Document(str(path))
    sec = doc.sections[0]
    paragraphs = [(p, p.text.strip()) for p in doc.paragraphs]
    texts = [text for _, text in paragraphs if text]
    first_index = next((i for i, (_, text) in enumerate(paragraphs) if text), None)

    blank_lines = 0
    if first_index is not None:
        for _, text in paragraphs[first_index + 1 :]:
            if text:
                break
            blank_lines += 1

    font_names = sorted({
        run.font.name
        for p, _ in paragraphs
        for run in p.runs
        if run.font.name
    })
    font_sizes = sorted({
        float(run.font.size.pt)
        for p, _ in paragraphs
        for run in p.runs
        if run.font.size
    })
    upper_text = "\n".join(texts).upper()
    sections_found = [
        section
        for section in profile.required_sections
        if section.upper() in upper_text
    ]

    return {
        "profile_id": profile.id,
        "page": {
            "width_cm": _round_cm(sec.page_width),
            "height_cm": _round_cm(sec.page_height),
        },
        "margins_cm": {
            "top": _round_cm(sec.top_margin),
            "left": _round_cm(sec.left_margin),
            "bottom": _round_cm(sec.bottom_margin),
            "right": _round_cm(sec.right_margin),
        },
        "paragraph_count": len(doc.paragraphs),
        "first_non_empty": texts[0] if texts else "",
        "blank_lines_after_header": blank_lines,
        "font_names": font_names,
        "font_sizes": font_sizes,
        "contains_oab": "OAB" in upper_text,
        "contains_local_data_hint": " DE 20" in upper_text,
        "required_sections_found": sections_found,
    }


def build_docx_report(path: Path, profile_id: str | None = None) -> dict[str, Any]:
    problemas = validar(path, profile_id=profile_id)
    return {
        "path": str(path),
        "status": "ok" if not problemas else "invalid_docx",
        "generated_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "problems": problemas,
        "structure": extract_docx_structure(path, profile_id),
    }


def build_run_report(
    *,
    profile: ValidationProfile,
    strict: bool,
    no_outbox: bool,
    summary: dict[str, int],
    items: list[dict[str, Any]],
) -> dict[str, Any]:
    return {
        "generated_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "profile": {
            "id": profile.id,
            "descricao": profile.descricao,
        },
        "strict": strict,
        "no_outbox": no_outbox,
        "summary": summary,
        "items": items,
    }


def write_json_report(path: Path, report: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")


def render_report_html(report: dict[str, Any]) -> str:
    """Converte um relatório JSON em HTML local para revisão humana."""
    profile = report.get("profile", {})
    summary = report.get("summary", {})
    items = report.get("items", [])
    rows = []
    for item in items:
        problems = item.get("problems") or []
        rows.append(
            "<tr>"
            f"<td>{escape(str(item.get('status', '')))}</td>"
            f"<td>{escape(str(item.get('docx') or ''))}</td>"
            f"<td>{escape(str(item.get('profile_id') or profile.get('id', '')))}</td>"
            f"<td>{escape('; '.join(map(str, problems)) or 'Sem violações formais')}</td>"
            "</tr>"
        )

    return f"""<!doctype html>
<html lang="pt-BR">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>Relatório de conformidade formal</title>
  <style>
    body {{ font-family: Arial, sans-serif; margin: 32px; color: #1f2933; }}
    h1 {{ margin-bottom: 4px; }}
    .muted {{ color: #667085; }}
    .warning {{ border-left: 4px solid #b54708; background: #fffaeb; padding: 12px; }}
    table {{ border-collapse: collapse; width: 100%; margin-top: 20px; }}
    th, td {{ border: 1px solid #d0d5dd; padding: 8px; text-align: left; vertical-align: top; }}
    th {{ background: #f2f4f7; }}
    code {{ background: #f2f4f7; padding: 2px 4px; border-radius: 4px; }}
  </style>
</head>
<body>
  <h1>Relatório de conformidade formal</h1>
  <p class="muted">Gerado em {escape(str(report.get('generated_at', '')))}</p>
  <p><strong>Perfil:</strong> {escape(str(profile.get('id', '')))} — {escape(str(profile.get('descricao', '')))}</p>
  <p><strong>Total:</strong> {escape(str(summary.get('total', 0)))} |
     <strong>Válidos:</strong> {escape(str(summary.get('validos', 0)))} |
     <strong>Bloqueados:</strong> {escape(str(summary.get('bloqueados', 0)))} |
     <strong>Falhas:</strong> {escape(str(summary.get('falhas', 0)))}</p>
  <div class="warning">
    Este relatório verifica conformidade formal automatizada. Ele não substitui revisão jurídica humana, conferência de documentos, assinatura, procuração ou regras locais de protocolo.
  </div>
  <table>
    <thead>
      <tr><th>Status</th><th>Documento</th><th>Perfil</th><th>Problemas</th></tr>
    </thead>
    <tbody>
      {''.join(rows) if rows else '<tr><td colspan="4">Nenhum item registrado.</td></tr>'}
    </tbody>
  </table>
</body>
</html>"""


def write_html_report(path: Path, report: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(render_report_html(report), encoding="utf-8")
