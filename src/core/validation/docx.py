"""Valida se um .docx atende às regras formais do projeto.

Retorna lista de violações (strings). Uma lista vazia significa que o documento
passou nas verificações determinísticas implementadas: página, margens, fonte,
alinhamentos, 7 linhas após o endereçamento, OAB, local/data, placeholders e
assinatura gráfica.

Não valida semântica jurídica — essa responsabilidade fica com o redator
(humano ou integrador externo) ao comparar o texto final contra o
`prompt_peticao.md`.
"""
from __future__ import annotations

import re
import sys
import unicodedata
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.core.profiles import ValidationProfile, get_profile

PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
TOP_MARGIN_CM = 3.0
LEFT_MARGIN_CM = 3.0
BOTTOM_MARGIN_CM = 2.0
RIGHT_MARGIN_CM = 2.0
FONTES_ACEITAS = {"Times New Roman", "Arial"}
TAM_FONTE_PT = 12
BLANKS_APOS_ENDERECO = 7
FIRST_LINE_INDENT_CM = 2.5

_EMU_POR_CM = 360000

BRACKET_PLACEHOLDER_RE = re.compile(r"\[[^\]]+\]")
PENDING_MARKER_RE = re.compile(
    r"\[(?:DADO\s+FALTANTE|REVISAR|INSERIR|PREENCHER|VERIFICAR|CONFIRMAR|TODO|FIXME)\b",
    re.IGNORECASE,
)
EXAMPLE_PLACEHOLDER_RE = re.compile(
    r"OAB/UF\s+0+|NOME DO (?:ADVOGADO|REQUERENTE|AUTOR|CLIENTE)",
    re.IGNORECASE,
)
DADO_FICTICIO_RE = re.compile(r"\b0{2,3}\.0{3}\.0{3}-0{2}\b|\b0{3}\.0{5}\.0{2}-0\b")
OAB_FORMAT_RE = re.compile(r"^\s*OAB\s*/\s*[A-Z]{2}\s*\d{1,3}(?:\.\d{3})*\s*$", re.IGNORECASE)
OAB_ANY_RE = re.compile(r"OAB", re.IGNORECASE)
LOCAL_DATA_RE = re.compile(
    r"^[A-ZÁÉÍÓÚÂÊÔÃÕÇ][\wÁÉÍÓÚÂÊÔÃÕÇáéíóúâêôãõç\s.-]+"
    r"(?:/[A-Z]{2})?,\s*\d{1,2}\s+de\s+\w+\s+de\s+\d{4}\.?\s*$",
    re.IGNORECASE,
)


def _cm(emu) -> float:
    return (emu or 0) / _EMU_POR_CM


def _normalize(value: str) -> str:
    nfkd = unicodedata.normalize("NFD", value)
    return "".join(c for c in nfkd if not unicodedata.combining(c)).upper()


def _header_matches(texto: str, profile: ValidationProfile) -> bool:
    primeira = next((linha.strip() for linha in texto.splitlines() if linha.strip()), "")
    primeira_norm = _normalize(primeira)
    return any(primeira_norm.startswith(_normalize(prefix)) for prefix in profile.header_prefixes)


def _has_disallowed_bracket_marker(texto: str, *, allow_pending_markers: bool) -> bool:
    for match in BRACKET_PLACEHOLDER_RE.finditer(texto):
        marker = match.group(0)
        if not allow_pending_markers or not PENDING_MARKER_RE.match(marker):
            return True
    return False


def validar_texto_protocolavel(
    texto: str,
    profile_id: str | None = None,
    *,
    allow_pending_markers: bool = False,
) -> list[str]:
    """Valida bloqueios formais antes de gerar/enfileirar uma peça."""
    profile = get_profile(profile_id)
    problemas: list[str] = []
    texto_limpo = texto.strip()
    texto_upper = texto_limpo.upper()

    if not texto_limpo:
        return ["texto da peça está vazio"]

    if EXAMPLE_PLACEHOLDER_RE.search(texto_limpo) or _has_disallowed_bracket_marker(
        texto_limpo,
        allow_pending_markers=allow_pending_markers,
    ):
        problemas.append("texto contém placeholders ou dados de exemplo")
    if DADO_FICTICIO_RE.search(texto_limpo):
        problemas.append("texto contém CPF/NIT fictício ou zerado")
    if not _header_matches(texto_limpo, profile):
        problemas.append(
            f"texto não inicia com endereçamento reconhecido pelo perfil {profile.id}"
        )
    if profile.require_oab and not any(
        OAB_FORMAT_RE.match(linha.strip()) for linha in texto_limpo.splitlines()
    ):
        problemas.append("texto não contém linha de OAB em formato reconhecido")
    if profile.require_local_data and not any(
        LOCAL_DATA_RE.match(linha.strip()) for linha in texto_limpo.splitlines()
    ):
        problemas.append("texto não contém local e data em formato reconhecido")
    if profile.require_oab and "TERMOS EM QUE" not in texto_upper and "NESTES TERMOS" not in texto_upper:
        problemas.append("texto não contém fechamento forense reconhecido")

    ausentes = [secao for secao in profile.required_sections if secao not in texto_upper]
    if ausentes:
        problemas.append(
            f"perfil {profile.id} sem seções mínimas: " + ", ".join(ausentes)
        )
    if profile.require_value_cause and "VALOR DA CAUSA" not in texto_upper:
        problemas.append(f"perfil {profile.id} exige valor da causa")

    return problemas


def validar(
    path: Path,
    profile_id: str | None = None,
    *,
    allow_pending_markers: bool = False,
) -> list[str]:
    profile = get_profile(profile_id)
    if not path.exists():
        return [f"arquivo não encontrado: {path}"]

    doc = Document(str(path))
    problemas: list[str] = []

    sec = doc.sections[0]
    for nome, atual, esperado in [
        ("largura da página", _cm(sec.page_width), PAGE_WIDTH_CM),
        ("altura da página", _cm(sec.page_height), PAGE_HEIGHT_CM),
    ]:
        if abs(atual - esperado) > 0.1:
            problemas.append(
                f"{nome}: {atual:.2f} cm (esperado A4: {esperado} cm)"
            )

    for nome, atual, esperado in [
        ("superior", _cm(sec.top_margin), TOP_MARGIN_CM),
        ("esquerda", _cm(sec.left_margin), LEFT_MARGIN_CM),
        ("inferior", _cm(sec.bottom_margin), BOTTOM_MARGIN_CM),
        ("direita", _cm(sec.right_margin), RIGHT_MARGIN_CM),
    ]:
        if abs(atual - esperado) > 0.1:
            problemas.append(
                f"margem {nome}: {atual:.2f} cm (esperado {esperado} cm)"
            )

    fontes_vistas: set[str] = set()
    tamanhos_vistos: set[int] = set()
    for p in doc.paragraphs:
        for run in p.runs:
            if run.font.name:
                fontes_vistas.add(run.font.name)
            if run.font.size:
                tamanhos_vistos.add(run.font.size.pt)
    fontes_invalidas = fontes_vistas - FONTES_ACEITAS
    if fontes_invalidas:
        problemas.append(
            f"fontes não autorizadas: {sorted(fontes_invalidas)} "
            f"(esperado: {sorted(FONTES_ACEITAS)})"
        )
    tamanhos_invalidos = {t for t in tamanhos_vistos if t != TAM_FONTE_PT}
    if tamanhos_invalidos:
        problemas.append(
            f"tamanhos de fonte não autorizados: {sorted(tamanhos_invalidos)} "
            f"(esperado: {TAM_FONTE_PT}pt)"
        )

    paragrafos = list(doc.paragraphs)
    paragrafos_texto = [(i, p, p.text.strip()) for i, p in enumerate(paragrafos)]
    primeiros_nao_vazios = [(i, p, t) for i, p, t in paragrafos_texto if t]

    if not primeiros_nao_vazios:
        problemas.append("documento vazio")
        return problemas

    texto_total = "\n".join(t for _, _, t in paragrafos_texto if t)
    problemas.extend(
        validar_texto_protocolavel(
            texto_total,
            profile.id,
            allow_pending_markers=allow_pending_markers,
        )
    )

    idx_p0, p0, t0 = primeiros_nao_vazios[0]
    t0_upper = t0.upper()
    if not _header_matches(t0, profile):
        problemas.append(
            f"1º parágrafo não parece endereçamento do perfil {profile.id}: {t0[:60]!r}"
        )
    if p0.alignment != WD_ALIGN_PARAGRAPH.CENTER:
        problemas.append("endereçamento não está centralizado")
    if not any(r.bold for r in p0.runs if r.text.strip()):
        problemas.append("endereçamento não está em negrito")

    brancos = 0
    for p in paragrafos[idx_p0 + 1 :]:
        if not p.text.strip():
            brancos += 1
        else:
            break
    if brancos < profile.min_blank_lines_after_header:
        problemas.append(
            f"{brancos} linhas em branco após endereçamento "
            f"(esperado {profile.min_blank_lines_after_header})"
        )

    ultimos = [t for _, _, t in paragrafos_texto if t][-4:]
    if profile.require_oab and not any("OAB" in t.upper() for t in ultimos):
        problemas.append("fechamento sem linha de OAB nos últimos parágrafos")

    paragrafo_oab = next(
        ((p, t) for _, p, t in reversed(paragrafos_texto) if OAB_ANY_RE.search(t)),
        None,
    )
    linha_oab = None
    if paragrafo_oab and profile.require_oab:
        p_oab, linha_oab = paragrafo_oab
        if not OAB_FORMAT_RE.match(linha_oab):
            problemas.append(f"linha OAB em formato não reconhecido: {linha_oab[:60]!r}")
        if p_oab.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            problemas.append("linha OAB não está centralizada")
        if not any(r.bold for r in p_oab.runs if r.text.strip()):
            problemas.append("linha OAB não está em negrito")

    if linha_oab and len(linha_oab) > 40:
        problemas.append(
            f"linha OAB muito longa (provavelmente nome+OAB juntos): {linha_oab[:60]!r}"
        )

    if profile.require_local_data and not any(
        LOCAL_DATA_RE.match(t) for _, _, t in paragrafos_texto if t
    ):
        problemas.append("local e data não encontrados em formato reconhecido")

    for _, p, t in paragrafos_texto:
        if not t or p.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
            continue
        indent = _cm(p.paragraph_format.first_line_indent)
        if abs(indent - FIRST_LINE_INDENT_CM) > 0.1:
            problemas.append(
                f"parágrafo sem recuo de {FIRST_LINE_INDENT_CM} cm: {t[:60]!r}"
            )
            break

    for _, _, t in paragrafos_texto:
        if "_" * 5 in t or "-" * 10 in t:
            problemas.append(
                "há linha de assinatura (traços/underscores) — proibido"
            )
            break

    return problemas


def main(argv: list[str]) -> int:
    if len(argv) not in {2, 4}:
        print("uso: python -m src.core.validation.docx <arquivo.docx> [--profile perfil]")
        return 2
    profile_id = None
    if len(argv) == 4:
        if argv[2] != "--profile":
            print("uso: python -m src.core.validation.docx <arquivo.docx> [--profile perfil]")
            return 2
        profile_id = argv[3]
    problemas = validar(Path(argv[1]), profile_id=profile_id)
    if not problemas:
        print("OK")
        return 0
    print(f"VIOLACOES ({len(problemas)}):")
    for v in problemas:
        print(f"  - {v}")
    return 1


if __name__ == "__main__":
    sys.exit(main(sys.argv))

