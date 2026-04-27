from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.core.validation.docx import validar, validar_texto_protocolavel
from src.infra.docx_render import renderizar


TEXTO_VALIDO = """EXCELENTÍSSIMO SENHOR DOUTOR JUIZ FEDERAL DA VARA DO JUIZADO ESPECIAL FEDERAL DA SUBSEÇÃO JUDICIÁRIA DE GOIÂNIA/GO

AÇÃO DE CONCESSÃO DE AUXÍLIO POR INCAPACIDADE TEMPORÁRIA

JOÃO DA SILVA, brasileiro, casado, pedreiro, portador do CPF 123.456.789-09, residente em Goiânia/GO, vem, respeitosamente, à presença de Vossa Excelência, propor a presente ação.

I - DOS FATOS

O requerente apresentou incapacidade laboral comprovada por documentos médicos.

II - DO DIREITO

Estão presentes os requisitos formais mínimos para análise supervisionada da peça.

III - DOS PEDIDOS

a) a concessão do benefício;
b) a produção de provas;
c) a condenação ao pagamento das parcelas vencidas.

IV - DO VALOR DA CAUSA

Dá-se à causa o valor de R$ 20.000,00.

Termos em que, pede deferimento.

Goiânia/GO, 6 de abril de 2026.

MARIA ADVOGADA DE EXEMPLO
OAB/GO 12.345"""


TEXTO_PROCURACAO = """PROCURAÇÃO AD JUDICIA ET EXTRA

OUTORGANTE: JOÃO DA SILVA, brasileiro, residente em Goiânia/GO.

OUTORGADA: MARIA ADVOGADA DE EXEMPLO, advogada inscrita na OAB/GO 12.345.

PODERES: representar o outorgante perante órgãos administrativos e judiciais, com poderes gerais para o foro, observados os limites da revisão humana obrigatória.

Goiânia/GO, 6 de abril de 2026.

JOÃO DA SILVA"""


def test_renderizar_docx_valido_com_acentos(tmp_path):
    destino = tmp_path / "peticao.docx"

    renderizar(TEXTO_VALIDO, destino)

    assert destino.exists()
    assert validar(destino) == []


def test_renderizar_define_pagina_a4(tmp_path):
    destino = tmp_path / "peticao.docx"
    renderizar(TEXTO_VALIDO, destino)

    doc = Document(destino)
    section = doc.sections[0]

    assert round(section.page_width.cm, 1) == 21.0
    assert round(section.page_height.cm, 1) == 29.7


def test_renderizar_negrito_pontual_em_qualificacao_e_labels(tmp_path):
    destino = tmp_path / "peticao.docx"
    renderizar(TEXTO_VALIDO, destino)

    doc = Document(destino)
    qualificacao = next(p for p in doc.paragraphs if p.text.startswith("JOÃO DA SILVA"))

    assert qualificacao.runs[0].text == "JOÃO DA SILVA"
    assert qualificacao.runs[0].bold is True


def test_renderizar_local_data_centralizado(tmp_path):
    destino = tmp_path / "peticao.docx"
    renderizar(TEXTO_VALIDO, destino)

    doc = Document(destino)
    local_data = next(p for p in doc.paragraphs if p.text.startswith("Goiânia/GO,"))

    assert local_data.alignment == WD_ALIGN_PARAGRAPH.CENTER


def test_validar_texto_protocolavel_bloqueia_placeholders():
    problemas = validar_texto_protocolavel(
        TEXTO_VALIDO.replace("JOÃO DA SILVA", "NOME DO REQUERENTE")
    )

    assert any("placeholders" in problema for problema in problemas)


def test_validar_docx_detecta_linha_de_assinatura(tmp_path):
    destino = tmp_path / "peticao.docx"
    renderizar(TEXTO_VALIDO + "\n\n__________", destino)

    problemas = validar(destino)

    assert any("linha de assinatura" in problema for problema in problemas)


def test_renderizar_procuracao_com_perfil_instrumento_mandato(tmp_path):
    destino = tmp_path / "procuracao.docx"

    renderizar(TEXTO_PROCURACAO, destino)

    assert destino.exists()
    assert validar(destino, profile_id="instrumento-mandato") == []
    assert validar_texto_protocolavel(TEXTO_PROCURACAO, "instrumento-mandato") == []
